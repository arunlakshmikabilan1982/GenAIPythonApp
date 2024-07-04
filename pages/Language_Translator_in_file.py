import os
import google.generativeai as genai
import streamlit as st
from io import BytesIO
import docx
from fpdf import FPDF

from Navigation import sidebar

sidebar()

api_key = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=api_key)

st.header("File Translator 🌐")
safetySettings = [ { 'category': 'HARM_CATEGORY_HATE_SPEECH', 'threshold': 'BLOCK_NONE' },
  {
    'category': 'HARM_CATEGORY_SEXUALLY_EXPLICIT',
    'threshold': 'BLOCK_NONE'
  },
  { 'category': 'HARM_CATEGORY_HARASSMENT', 'threshold': 'BLOCK_NONE' },
  {
    'category': 'HARM_CATEGORY_DANGEROUS_CONTENT',
    'threshold': 'BLOCK_NONE'
  }
  ]

model = genai.GenerativeModel('gemini-pro')

def translate_text(text, source_lang, target_lang):
    response = model.generate_content(f"Translate the following sentence from {source_lang} to {target_lang}: {text}", safety_settings=safetySettings)
    return response.text

def translate_file_in_docx(file, source_lang, target_lang):
    doc = docx.Document(BytesIO(file.read()))
    translated_doc = docx.Document()
    with st.spinner('Translating DOCX...'):
        for para in doc.paragraphs:
            text = para.text
            if text:
                translated_text = translate_text(text, source_lang, target_lang)
                translated_doc.add_paragraph(translated_text)
    translated_doc_bytes = BytesIO()
    translated_doc.save(translated_doc_bytes)
    translated_doc_bytes.seek(0)
    return translated_doc_bytes

# def translate_file_in_pdf(file, source_lang, target_lang, file_name):
#     doc = docx.Document(BytesIO(file.read()))
#     fpdf = FPDF()
#     fpdf.add_page()    
#     fpdf.set_font("Arial", size=12)
#      with st.spinner('Translating PDF...'):
#       for para in doc.paragraphs:
#            text = para.text
#            if text:
#                translated_text = translate_text(text, source_lang, target_lang)
#                translated_text = translated_text.encode('latin-1', 'replace').decode('latin-1')
#                fpdf.multi_cell(0, 10, txt=translated_text)
#             fpdf.multi_cell(0, 10, txt=translated_text)
#     translated_pdf_filename = "Translated_file.pdf"
#     fpdf.output(translated_pdf_filename)
#     return translated_pdf_filename

def main():
    # st.set_page_config(page_title="File Translator", page_icon="📄")
    # st.header("File Translator 🌐")

    uploaded_file = st.file_uploader("Upload a file (.docx, .pdf, .txt)", type=['docx', 'pdf', 'txt'])
    if uploaded_file is not None:
        source_language = st.selectbox("Select Source Language", ["English", "French", "Spanish", "German", "Arabic"])
        target_language = st.selectbox("Select Target Language", ["French", "Chinese", "Arabic", "Spanish", "German"])

        # docx_btn, pdf_btn = st.columns([1, 1])
        # with docx_btn:
        if st.button("Translate into DOCX"):
            translated_docx = translate_file_in_docx(uploaded_file, source_language.lower(), target_language.lower())
            st.write("Translation complete!")
            st.download_button(
                label="Download Translated DOCX",
                data=translated_docx,
                file_name=f"Translated_{uploaded_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        # with pdf_btn:
        #     if st.button("Translate into PDF"):
        #         translated_pdf = translate_file_in_pdf(uploaded_file, source_language.lower(), target_language.lower(), uploaded_file.name)
        #         st.write("Translation complete!")
        #         st.download_button(
        #             label="Download Translated PDF",
        #             data=translated_pdf,
        #             file_name="Translated_file.pdf",
        #             mime="application/pdf"
        #         )

if __name__ == "__main__":
    main()
